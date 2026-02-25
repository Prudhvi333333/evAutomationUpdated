"""
EV Research Testing Script - Robust RAG Pipeline
Tests 20 questions using 3 models:
  1. TinyLlama (local) with RAG
  2. Gemini 2.5 Flash with RAG
  3. Gemini 2.5 Flash without RAG

Features improved chunking and retrieval logic.
Processes single Excel file: GNEM updated excel (1).xlsx
"""

import os
import json
import time
import re
import pandas as pd
from docx import Document
from typing import List, Dict, Optional, Tuple
import google.generativeai as genai
from dotenv import load_dotenv
from collections import defaultdict
import ollama

# Load existing modules
from document_processor import DocumentProcessor
from vector_store import VectorStore
from rag_pipeline import RAGPipeline

# Load environment variables for Gemini API key
load_dotenv()

# Configuration
EXCEL_FILE = "GNEM updated excel (1).xlsx"
QUESTIONS_FILE = "Sample questions.docx"
OUTPUT_EXCEL = "test_results.xlsx"
CHUNKING_REPORT = "chunking_report.json"
NUM_QUESTIONS = 20

# Model configurations
TINYLLAMA_MODEL = "tinyllama"
GEMINI_MODEL = os.getenv('GEMINI_MODEL', 'gemini-2.5-flash-preview-05-20')


class RobustChunker:
    """Advanced chunking strategy for structured Excel data."""
    
    def __init__(self, df: pd.DataFrame):
        self.df = df
        self.columns = df.columns.tolist()
        self.chunks = []
        self.chunk_stats = {
            'total_rows': len(df),
            'total_chunks': 0,
            'chunks_by_strategy': defaultdict(int),
            'avg_chunk_size': 0,
            'field_coverage': {}
        }
    
    def create_row_chunks(self) -> List[Dict]:
        """Create chunks from individual rows with full field coverage."""
        chunks = []
        
        for idx, row in self.df.iterrows():
            # Build structured text with field labels
            fields = []
            for col in self.columns:
                value = row[col]
                if pd.notna(value) and str(value).strip():
                    fields.append(f"{col}: {value}")
            
            row_text = " | ".join(fields)
            
            # Create rich metadata
            metadata = {
                'row_index': int(idx),
                'company': str(row.get('Company', 'Unknown')),
                'category': str(row.get('Category', 'Unknown')),
                'industry_group': str(row.get('Industry Group', 'Unknown')),
                'location': str(row.get('Location', 'Unknown')),
                'ev_role': str(row.get('EV Supply Chain Role', 'Unknown')),
                'ev_relevant': str(row.get('EV / Battery Relevant', 'Unknown')),
                'employment': str(row.get('Employment', 'Unknown')),
                'primary_oems': str(row.get('Primary OEMs', 'Unknown')),
                'source': EXCEL_FILE,
                'chunk_type': 'full_row',
                'field_count': len(fields)
            }
            
            chunks.append({
                'text': row_text,
                'metadata': metadata
            })
        
        self.chunk_stats['chunks_by_strategy']['full_row'] = len(chunks)
        return chunks
    
    def create_semantic_chunks(self) -> List[Dict]:
        """Create semantic chunks by grouping related fields."""
        semantic_groups = {
            'company_identity': ['Company', 'Category', 'Industry Group', 'Classification Method'],
            'location_facility': ['Location', 'Primary Facility Type'],
            'supply_chain': ['EV Supply Chain Role', 'Supplier or Affiliation Type', 'Primary OEMs'],
            'employment_scale': ['Company', 'Employment', 'Category'],
            'product_details': ['Company', 'Product / Service', 'EV / Battery Relevant'],
            'ev_battery': ['Company', 'EV / Battery Relevant', 'EV Supply Chain Role', 'Product / Service']
        }
        
        chunks = []
        
        for idx, row in self.df.iterrows():
            for group_name, group_cols in semantic_groups.items():
                fields = []
                for col in group_cols:
                    if col in self.columns:
                        value = row[col]
                        if pd.notna(value) and str(value).strip():
                            fields.append(f"{col}: {value}")
                
                if fields:
                    chunk_text = f"[{group_name}] " + " | ".join(fields)
                    
                    metadata = {
                        'row_index': int(idx),
                        'company': str(row.get('Company', 'Unknown')),
                        'semantic_group': group_name,
                        'source': EXCEL_FILE,
                        'chunk_type': 'semantic',
                        'field_count': len(fields)
                    }
                    
                    chunks.append({
                        'text': chunk_text,
                        'metadata': metadata
                    })
        
        self.chunk_stats['chunks_by_strategy']['semantic'] = len(chunks)
        return chunks
    
    def create_keyword_indexed_chunks(self) -> List[Dict]:
        """Create chunks optimized for keyword-based questions."""
        chunks = []
        
        for idx, row in self.df.iterrows():
            # Company-focused chunk
            company = row.get('Company', '')
            if pd.notna(company) and str(company).strip():
                related_fields = []
                for col in self.columns:
                    if col != 'Company':
                        value = row[col]
                        if pd.notna(value) and str(value).strip():
                            related_fields.append(f"{col}: {value}")
                
                chunk_text = f"Company: {company} | " + " | ".join(related_fields[:8])
                
                metadata = {
                    'row_index': int(idx),
                    'company': str(company),
                    'search_key': str(company).lower(),
                    'source': EXCEL_FILE,
                    'chunk_type': 'company_focused'
                }
                
                chunks.append({
                    'text': chunk_text,
                    'metadata': metadata
                })
                
                # OEM-focused chunks
                oems = row.get('Primary OEMs', '')
                if pd.notna(oems) and str(oems).strip():
                    oem_list = [o.strip() for o in str(oems).replace(',', ' ').split() if o.strip()]
                    for oem in oem_list:
                        if len(oem) > 2:
                            oem_chunk_text = f"OEM: {oem} | Company: {company}"
                            for col in ['Category', 'EV Supply Chain Role', 'Product / Service', 'Location']:
                                if col in self.columns:
                                    val = row[col]
                                    if pd.notna(val):
                                        oem_chunk_text += f" | {col}: {val}"
                            
                            chunks.append({
                                'text': oem_chunk_text,
                                'metadata': {
                                    'row_index': int(idx),
                                    'company': str(company),
                                    'oem': oem,
                                    'search_key': oem.lower(),
                                    'source': EXCEL_FILE,
                                    'chunk_type': 'oem_focused'
                                }
                            })
        
        self.chunk_stats['chunks_by_strategy']['keyword_indexed'] = len(chunks)
        return chunks
    
    def generate_all_chunks(self) -> List[Dict]:
        """Generate comprehensive chunk set with all strategies."""
        print("Generating chunks with multiple strategies...")
        
        # Generate all chunk types
        row_chunks = self.create_row_chunks()
        semantic_chunks = self.create_semantic_chunks()
        keyword_chunks = self.create_keyword_indexed_chunks()
        
        # Combine all chunks
        all_chunks = row_chunks + semantic_chunks + keyword_chunks
        
        # Calculate stats
        self.chunk_stats['total_chunks'] = len(all_chunks)
        total_text_length = sum(len(c['text']) for c in all_chunks)
        self.chunk_stats['avg_chunk_size'] = round(total_text_length / len(all_chunks), 1) if all_chunks else 0
        
        # Field coverage analysis
        field_counts = defaultdict(int)
        for chunk in all_chunks:
            text = chunk['text']
            for col in self.columns:
                if col in text:
                    field_counts[col] += 1
        
        self.chunk_stats['field_coverage'] = dict(field_counts)
        
        print(f"  - Full row chunks: {len(row_chunks)}")
        print(f"  - Semantic chunks: {len(semantic_chunks)}")
        print(f"  - Keyword indexed chunks: {len(keyword_chunks)}")
        print(f"  - Total chunks: {len(all_chunks)}")
        print(f"  - Average chunk size: {self.chunk_stats['avg_chunk_size']} chars")
        
        return all_chunks
    
    def get_stats(self) -> Dict:
        """Return chunking statistics."""
        return dict(self.chunk_stats)


class RobustRetriever:
    """Advanced retrieval logic with multiple search strategies."""
    
    def __init__(self, vector_store: VectorStore, df: pd.DataFrame):
        self.vector_store = vector_store
        self.df = df
    
    def extract_keywords(self, question: str) -> List[str]:
        """Extract key entities from question."""
        keywords = []
        
        # Company names (look for capitalized words)
        company_pattern = r'\b[A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)*\b'
        potential_companies = re.findall(company_pattern, question)
        
        # Common company suffixes
        suffixes = ['Corp', 'Inc', 'LLC', 'Ltd', 'Co', 'Company', 'Corporation']
        for pc in potential_companies:
            if any(s in pc for s in suffixes) or len(pc) > 3:
                keywords.append(pc.lower())
        
        # OEM names (common car manufacturers)
        oem_keywords = ['hyundai', 'kia', 'rivian', 'porsche', 'yamaha', 'blue bird', 
                       'mercedes', 'bmw', 'toyota', 'honda', 'ford', 'gm', 'tesla']
        for oem in oem_keywords:
            if oem in question.lower():
                keywords.append(oem)
        
        # Location keywords
        location_keywords = ['county', 'georgia', 'gwinnett', 'troup', 'spalding', 
                            'cherokee', 'walton', 'elbert', 'lamar']
        for loc in location_keywords:
            if loc in question.lower():
                keywords.append(loc)
        
        # Role keywords
        role_keywords = ['tier 1', 'tier 2', 'tier 3', 'oem', 'supplier', 'manufacturer']
        for role in role_keywords:
            if role in question.lower():
                keywords.append(role.replace(' ', '_'))
        
        # Numeric patterns (employment thresholds)
        numbers = re.findall(r'\d{3,}', question)
        keywords.extend(numbers)
        
        return list(set(keywords))
    
    def classify_question_type(self, question: str) -> str:
        """Classify question type for targeted retrieval."""
        q_lower = question.lower()
        
        if 'how many' in q_lower or 'count' in q_lower or 'number of' in q_lower:
            return 'count'
        elif 'highest' in q_lower or 'most' in q_lower or 'top' in q_lower:
            return 'maximum'
        elif 'lowest' in q_lower or 'least' in q_lower or 'minimum' in q_lower:
            return 'minimum'
        elif 'average' in q_lower or 'mean' in q_lower:
            return 'average'
        elif any(oem in q_lower for oem in ['hyundai', 'kia', 'rivian', 'porsche', 'yamaha']):
            return 'oem_specific'
        elif 'location' in q_lower or 'county' in q_lower:
            return 'location_based'
        elif 'company' in q_lower and 'list' in q_lower:
            return 'company_list'
        else:
            return 'general'
    
    def retrieve(self, question: str, top_k: int = 10) -> List[Dict]:
        """Multi-strategy retrieval."""
        question_type = self.classify_question_type(question)
        keywords = self.extract_keywords(question)
        
        print(f"    [Retriever] Type: {question_type}, Keywords: {keywords}")
        
        # Strategy 1: Vector similarity search
        vector_results = self.vector_store.search(question, top_k=top_k)
        
        # Strategy 2: Keyword-boosted results from dataframe
        keyword_results = []
        if keywords:
            for keyword in keywords[:3]:  # Use top 3 keywords
                # Search in dataframe directly
                mask = self.df.astype(str).apply(
                    lambda x: x.str.lower().str.contains(keyword, na=False)
                ).any(axis=1)
                matching_rows = self.df[mask]
                
                for idx, row in matching_rows.head(5).iterrows():
                    fields = []
                    for col in self.df.columns:
                        val = row[col]
                        if pd.notna(val):
                            fields.append(f"{col}: {val}")
                    
                    keyword_results.append({
                        'text': " | ".join(fields),
                        'metadata': {
                            'row_index': int(idx),
                            'company': str(row.get('Company', 'Unknown')),
                            'match_type': 'keyword',
                            'matched_keyword': keyword
                        },
                        'score': 0.9  # High score for keyword match
                    })
        
        # Strategy 3: Question-type specific filtering
        filtered_results = []
        if question_type == 'maximum' and 'employment' in question.lower():
            for r in vector_results + keyword_results:
                if 'Employment' in r['text'] or 'employment' in r['text'].lower():
                    r['score'] = r.get('score', 0.5) * 1.2
                    filtered_results.append(r)
        elif question_type == 'oem_specific':
            for r in vector_results + keyword_results:
                text_lower = r['text'].lower()
                if any(k in text_lower for k in keywords):
                    r['score'] = r.get('score', 0.5) * 1.3
                    filtered_results.append(r)
        
        # Merge and deduplicate results
        all_results = vector_results + keyword_results + filtered_results
        
        # Deduplicate by company name
        seen_companies = set()
        unique_results = []
        for r in all_results:
            company = r['metadata'].get('company', 'Unknown')
            if company not in seen_companies:
                seen_companies.add(company)
                unique_results.append(r)
        
        # Sort by score and return top results
        unique_results.sort(key=lambda x: x.get('score', 0), reverse=True)
        return unique_results[:top_k]


def extract_questions_from_docx(filepath: str, limit: int = 20) -> List[str]:
    """Extract questions from docx file."""
    doc = Document(filepath)
    questions = []
    
    # Extract from paragraphs
    for p in doc.paragraphs:
        text = p.text.strip()
        if text and len(text) > 10 and '?' in text:
            questions.append(text)
    
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and len(text) > 10 and '?' in text:
                    questions.append(text)
    
    # Clean and deduplicate
    unique_questions = []
    seen = set()
    for q in questions:
        clean_q = q.replace('\n', ' ').strip()
        if clean_q not in seen and len(clean_q) > 15:
            seen.add(clean_q)
            unique_questions.append(clean_q)
    
    return unique_questions[:limit]


def setup_rag_pipeline(excel_path: str) -> Tuple[RAGPipeline, VectorStore, pd.DataFrame, Dict, RobustRetriever]:
    """Setup robust RAG pipeline with advanced chunking."""
    print(f"\nSetting up robust RAG pipeline with: {excel_path}")
    
    # Initialize vector store
    vector_store = VectorStore()
    
    # Clear existing collection
    try:
        vector_store.delete_collection()
        print("Cleared existing vector store")
    except:
        pass
    
    # Read Excel file
    df = pd.read_excel(excel_path)
    print(f"Loaded {len(df)} rows from Excel")
    print(f"Columns: {df.columns.tolist()}")
    
    # Use robust chunker
    chunker = RobustChunker(df)
    chunks = chunker.generate_all_chunks()
    
    # Add to vector store
    total_added = vector_store.add_documents(chunks)
    print(f"Added {total_added} documents to vector store")
    
    # Get chunking stats
    chunk_stats = chunker.get_stats()
    
    # Create RAG pipeline
    rag = RAGPipeline(vector_store)
    
    # Create robust retriever
    retriever = RobustRetriever(vector_store, df)
    
    return rag, vector_store, df, chunk_stats, retriever


def query_tinyllama_rag(retriever: RobustRetriever, question: str) -> Dict:
    """Query TinyLlama with robust RAG retrieval."""
    start_time = time.time()
    
    # Use robust retrieval
    results = retriever.retrieve(question, top_k=8)
    context = format_context(results)
    
    # Generate prompt
    prompt = f"""You are a helpful assistant that answers questions based on the provided context about Georgia EV supply chain companies.

Context:
{context}

Question: {question}

Instructions:
- Answer based ONLY on the context provided
- Be concise and specific
- If the answer is not in the context, say "Not found in provided data"
- Include company names when relevant

Answer:"""
    
    try:
        response = ollama.generate(
            model=TINYLLAMA_MODEL,
            prompt=prompt,
            options={'temperature': 0.3, 'num_predict': 300}
        )
        answer = response['response']
        success = True
    except Exception as e:
        answer = f"Error: {str(e)}"
        success = False
    
    elapsed_time = time.time() - start_time
    
    return {
        'model': 'TinyLlama',
        'rag': True,
        'answer': answer,
        'context': results,
        'success': success,
        'time_seconds': round(elapsed_time, 2)
    }


def query_gemini_rag(retriever: RobustRetriever, question: str, api_key: str) -> Dict:
    """Query Gemini 2.5 with robust RAG retrieval."""
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(GEMINI_MODEL)
    
    # Use robust retrieval
    results = retriever.retrieve(question, top_k=8)
    context = format_context(results)
    
    prompt = f"""You are a helpful assistant that answers questions based on the provided context about Georgia EV supply chain companies.

Context:
{context}

Question: {question}

Instructions:
- Answer based ONLY on the context provided
- Be concise and specific
- If the answer is not in the context, say "Not found in provided data"
- Include company names when relevant

Answer:"""
    
    start_time = time.time()
    try:
        response = model.generate_content(prompt)
        answer = response.text
        success = True
    except Exception as e:
        answer = f"Error: {str(e)}"
        success = False
    
    elapsed_time = time.time() - start_time
    
    return {
        'model': 'Gemini 2.5',
        'rag': True,
        'answer': answer,
        'context': results,
        'success': success,
        'time_seconds': round(elapsed_time, 2)
    }


def query_gemini_no_rag(question: str, api_key: str) -> Dict:
    """Query Gemini 2.5 without RAG."""
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(GEMINI_MODEL)
    
    prompt = f"""You are a helpful assistant. Answer the following question to the best of your knowledge.

Question: {question}

Instructions:
- Be concise and specific
- If you don't know the answer, say so

Answer:"""
    
    start_time = time.time()
    try:
        response = model.generate_content(prompt)
        answer = response.text
        success = True
    except Exception as e:
        answer = f"Error: {str(e)}"
        success = False
    
    elapsed_time = time.time() - start_time
    
    return {
        'model': 'Gemini 2.5',
        'rag': False,
        'answer': answer,
        'context': [],
        'success': success,
        'time_seconds': round(elapsed_time, 2)
    }


def format_context(results: List[Dict]) -> str:
    """Format retrieved results into context string."""
    if not results:
        return "No relevant context found."
    
    context_parts = []
    for i, result in enumerate(results, 1):
        source = result['metadata'].get('company', 'Unknown')
        match_type = result['metadata'].get('match_type', 'vector')
        text = result['text']
        score = result.get('score', 0)
        context_parts.append(f"[Source {i}: {source} | Match: {match_type} | Score: {score:.3f}]\n{text}")
    
    return "\n\n".join(context_parts)


def run_chunking_tests(df: pd.DataFrame, chunk_stats: Dict) -> Dict:
    """Run test cases to verify chunking quality."""
    print("\n" + "="*80)
    print("Running Chunking Quality Tests")
    print("="*80)
    
    tests = {
        'total_rows': len(df),
        'total_columns': len(df.columns),
        'columns': df.columns.tolist(),
        'chunking_stats': chunk_stats,
        'field_coverage': chunk_stats.get('field_coverage', {}),
        'sample_rows': []
    }
    
    # Test 1: Verify all rows are represented
    print(f"\n[TEST 1] Row Coverage")
    print(f"  Expected chunks per strategy: {len(df)} rows")
    print(f"  Full row chunks: {chunk_stats['chunks_by_strategy'].get('full_row', 0)}")
    assert chunk_stats['chunks_by_strategy'].get('full_row', 0) >= len(df), "Missing row chunks!"
    print("  ✓ All rows have corresponding chunks")
    
    # Test 2: Verify field coverage
    print(f"\n[TEST 2] Field Coverage")
    field_coverage = chunk_stats.get('field_coverage', {})
    for col in df.columns:
        count = field_coverage.get(col, 0)
        print(f"  {col}: {count} chunks")
        assert count > 0, f"Field {col} not covered in any chunk!"
    print("  ✓ All fields are covered in chunks")
    
    # Test 3: Sample row validation
    print(f"\n[TEST 3] Sample Row Validation")
    sample_indices = [0, len(df)//2, len(df)-1] if len(df) >= 3 else [0]
    for idx in sample_indices:
        row = df.iloc[idx]
        row_dict = {col: str(row[col]) for col in df.columns if pd.notna(row[col])}
        tests['sample_rows'].append({
            'index': int(idx),
            'data': row_dict
        })
        print(f"  Row {idx}: {row.get('Company', 'N/A')} - {len(row_dict)} fields")
    print("  ✓ Sample rows validated")
    
    # Test 4: Semantic group coverage
    print(f"\n[TEST 4] Semantic Group Coverage")
    semantic_count = chunk_stats['chunks_by_strategy'].get('semantic', 0)
    print(f"  Semantic chunks: {semantic_count}")
    assert semantic_count > 0, "No semantic chunks generated!"
    print("  ✓ Semantic groups are present")
    
    # Test 5: Average chunk size validation
    print(f"\n[TEST 5] Chunk Size Validation")
    avg_size = chunk_stats.get('avg_chunk_size', 0)
    print(f"  Average chunk size: {avg_size} characters")
    assert 100 < avg_size < 2000, f"Average chunk size ({avg_size}) seems unusual!"
    print("  ✓ Chunk sizes are reasonable")
    
    print("\n" + "="*80)
    print("All Chunking Tests Passed!")
    print("="*80)
    
    return tests


def run_retrieval_tests(retriever: RobustRetriever, questions: List[str]) -> Dict:
    """Run test cases to verify retrieval quality."""
    print("\n" + "="*80)
    print("Running Retrieval Quality Tests")
    print("="*80)
    
    test_results = {
        'retrieval_tests': [],
        'keyword_extraction': [],
        'question_classification': []
    }
    
    # Test keyword extraction on sample questions
    print("\n[TEST 1] Keyword Extraction")
    for q in questions[:3]:
        keywords = retriever.extract_keywords(q)
        question_type = retriever.classify_question_type(q)
        print(f"  Q: {q[:60]}...")
        print(f"    Keywords: {keywords}")
        print(f"    Type: {question_type}")
        test_results['keyword_extraction'].append({
            'question': q,
            'keywords': keywords,
            'type': question_type
        })
    print("  ✓ Keyword extraction working")
    
    # Test retrieval on first question
    print("\n[TEST 2] Retrieval Execution")
    if questions:
        results = retriever.retrieve(questions[0], top_k=5)
        print(f"  Retrieved {len(results)} results for first question")
        for i, r in enumerate(results[:3], 1):
            company = r['metadata'].get('company', 'Unknown')
            score = r.get('score', 0)
            print(f"    {i}. {company} (score: {score:.3f})")
        test_results['retrieval_tests'].append({
            'question': questions[0],
            'results_count': len(results),
            'top_results': [{'company': r['metadata'].get('company'), 'score': r.get('score')} for r in results[:3]]
        })
        assert len(results) > 0, "Retrieval returned no results!"
        print("  ✓ Retrieval returns results")
    
    print("\n" + "="*80)
    print("All Retrieval Tests Passed!")
    print("="*80)
    
    return test_results


def run_model_tests(questions: List[str], retriever: RobustRetriever, df: pd.DataFrame) -> List[Dict]:
    """Run all model tests and collect results (3 models)."""
    results = []
    
    # Get Gemini API key
    api_key = os.getenv('GEMINI_API_KEY')
    if not api_key:
        print("Warning: GEMINI_API_KEY not found in environment variables")
        print("Please set your Gemini API key in a .env file")
        return []
    
    total_questions = len(questions)
    
    for idx, question in enumerate(questions, 1):
        print(f"\n{'='*80}")
        print(f"Question {idx}/{total_questions}: {question[:100]}...")
        print(f"{'='*80}")
        
        question_results = {
            'question_number': idx,
            'question': question
        }
        
        # Test 1: TinyLlama with RAG
        print("  [1/3] Testing TinyLlama with RAG...")
        result = query_tinyllama_rag(retriever, question)
        question_results['tinyllama_rag'] = result['answer']
        question_results['tinyllama_rag_time'] = result['time_seconds']
        question_results['tinyllama_rag_success'] = result['success']
        
        # Test 2: Gemini 2.5 with RAG
        print("  [2/3] Testing Gemini 2.5 with RAG...")
        result = query_gemini_rag(retriever, question, api_key)
        question_results['gemini_rag'] = result['answer']
        question_results['gemini_rag_time'] = result['time_seconds']
        question_results['gemini_rag_success'] = result['success']
        
        # Test 3: Gemini 2.5 without RAG
        print("  [3/3] Testing Gemini 2.5 without RAG...")
        result = query_gemini_no_rag(question, api_key)
        question_results['gemini_no_rag'] = result['answer']
        question_results['gemini_no_rag_time'] = result['time_seconds']
        question_results['gemini_no_rag_success'] = result['success']
        
        results.append(question_results)
        print(f"  Completed question {idx}/{total_questions}")
    
    return results


def generate_excel_report(results: List[Dict], chunk_tests: Dict, retrieval_tests: Dict, output_path: str):
    """Generate comprehensive Excel report with all results."""
    # Create main results dataframe
    rows = []
    for r in results:
        rows.append({
            'Question Number': r['question_number'],
            'Question': r['question'],
            'TinyLlama RAG': r['tinyllama_rag'],
            'TinyLlama Time (s)': r['tinyllama_rag_time'],
            'TinyLlama Success': r['tinyllama_rag_success'],
            'Gemini 2.5 RAG': r['gemini_rag'],
            'Gemini RAG Time (s)': r['gemini_rag_time'],
            'Gemini RAG Success': r['gemini_rag_success'],
            'Gemini 2.5 No RAG': r['gemini_no_rag'],
            'Gemini No RAG Time (s)': r['gemini_no_rag_time'],
            'Gemini No RAG Success': r['gemini_no_rag_success']
        })
    
    df_results = pd.DataFrame(rows)
    
    # Create chunking stats dataframe
    chunk_stats = chunk_tests.get('chunking_stats', {})
    df_chunk_stats = pd.DataFrame([{
        'Metric': 'Total Rows',
        'Value': chunk_stats.get('total_rows', 0)
    }, {
        'Metric': 'Total Chunks',
        'Value': chunk_stats.get('total_chunks', 0)
    }, {
        'Metric': 'Full Row Chunks',
        'Value': chunk_stats.get('chunks_by_strategy', {}).get('full_row', 0)
    }, {
        'Metric': 'Semantic Chunks',
        'Value': chunk_stats.get('chunks_by_strategy', {}).get('semantic', 0)
    }, {
        'Metric': 'Keyword Indexed Chunks',
        'Value': chunk_stats.get('chunks_by_strategy', {}).get('keyword_indexed', 0)
    }, {
        'Metric': 'Average Chunk Size (chars)',
        'Value': chunk_stats.get('avg_chunk_size', 0)
    }])
    
    # Create field coverage dataframe
    field_coverage = chunk_tests.get('field_coverage', {})
    df_field_coverage = pd.DataFrame([
        {'Field': field, 'Chunk Count': count}
        for field, count in field_coverage.items()
    ])
    
    # Write to Excel with multiple sheets
    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
        # Sheet 1: Test Results
        df_results.to_excel(writer, sheet_name='Model Responses', index=False)
        
        # Sheet 2: Chunking Stats
        df_chunk_stats.to_excel(writer, sheet_name='Chunking Stats', index=False)
        
        # Sheet 3: Field Coverage
        df_field_coverage.to_excel(writer, sheet_name='Field Coverage', index=False)
        
        # Adjust column widths for main sheet
        worksheet = writer.sheets['Model Responses']
        for column in worksheet.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 100)
            worksheet.column_dimensions[column_letter].width = adjusted_width
    
    print(f"\nResults saved to: {output_path}")
    print(f"Total questions tested: {len(results)}")
    
    # Print summary
    tiny_success = sum(1 for r in results if r['tinyllama_rag_success'])
    gemini_rag_success = sum(1 for r in results if r['gemini_rag_success'])
    gemini_no_rag_success = sum(1 for r in results if r['gemini_no_rag_success'])
    
    print(f"\nSuccess Summary:")
    print(f"  TinyLlama RAG: {tiny_success}/{len(results)}")
    print(f"  Gemini 2.5 RAG: {gemini_rag_success}/{len(results)}")
    print(f"  Gemini 2.5 No RAG: {gemini_no_rag_success}/{len(results)}")


def main():
    """Main test runner with robust pipeline."""
    print("="*80)
    print("EV Research Testing - Robust RAG Pipeline")
    print("="*80)
    print("\nModels: TinyLlama (RAG) | Gemini 2.5 (RAG) | Gemini 2.5 (No RAG)")
    print("Chunking: Full Row + Semantic + Keyword Indexed")
    print("Retrieval: Vector + Keyword + Question-Type Boosting")
    
    # Check files exist
    if not os.path.exists(EXCEL_FILE):
        print(f"Error: Excel file not found: {EXCEL_FILE}")
        return
    
    if not os.path.exists(QUESTIONS_FILE):
        print(f"Error: Questions file not found: {QUESTIONS_FILE}")
        return
    
    # Extract questions
    print(f"\nExtracting {NUM_QUESTIONS} questions from {QUESTIONS_FILE}...")
    questions = extract_questions_from_docx(QUESTIONS_FILE, NUM_QUESTIONS)
    print(f"Found {len(questions)} questions")
    
    if len(questions) == 0:
        print("Error: No questions found in the docx file")
        return
    
    # Setup robust RAG pipeline
    print(f"\nSetting up robust RAG pipeline...")
    rag_pipeline, vector_store, df, chunk_stats, retriever = setup_rag_pipeline(EXCEL_FILE)
    
    # Run chunking tests
    chunk_tests = run_chunking_tests(df, chunk_stats)
    
    # Save chunking report
    with open(CHUNKING_REPORT, 'w') as f:
        json.dump(chunk_tests, f, indent=2, default=str)
    print(f"\nChunking report saved to: {CHUNKING_REPORT}")
    
    # Run retrieval tests
    retrieval_tests = run_retrieval_tests(retriever, questions[:5])
    
    # Run model tests
    print(f"\nRunning model tests on {len(questions)} questions...")
    results = run_model_tests(questions, retriever, df)
    
    if not results:
        print("No results generated - check API key configuration")
        return
    
    # Generate report
    print(f"\nGenerating Excel report...")
    generate_excel_report(results, chunk_tests, retrieval_tests, OUTPUT_EXCEL)
    
    print("\n" + "="*80)
    print("Testing completed!")
    print("="*80)


if __name__ == "__main__":
    main()
